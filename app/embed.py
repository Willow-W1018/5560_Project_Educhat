"""
Enhanced Text Embedding Module

Features:
- Support for multiple file formats: PDF, TXT, DOCX, MD, HTML
- Configurable chunk size and overlap
- Text preprocessing and cleaning
"""

import PyPDF2
from sentence_transformers import SentenceTransformer
import os
import re
from typing import List, Tuple, Optional
from app.db import vector_db

# Load the sentence transformer model
# Try local model first, fallback to download
MODEL_PATH = 'Models/paraphrase-MiniLM-L6-v2'
if os.path.exists(MODEL_PATH):
    model = SentenceTransformer(MODEL_PATH)
else:
    model = SentenceTransformer('paraphrase-MiniLM-L6-v2')


def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    pdf_reader = PyPDF2.PdfReader(file.file)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def extract_text_from_txt(file) -> str:
    """Extract text from TXT file"""
    content = file.file.read()
    # Try different encodings
    for encoding in ['utf-8', 'latin-1', 'cp1252']:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode('utf-8', errors='ignore')


def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        # Save to temp file
        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(file.file.read())
            tmp_path = tmp.name
        
        doc = Document(tmp_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up temp file
        os.unlink(tmp_path)
        return text
    except ImportError:
        raise ValueError("python-docx is required for DOCX files. Install with: pip install python-docx")


def extract_text_from_markdown(file) -> str:
    """Extract text from Markdown file, removing markdown syntax"""
    content = file.file.read().decode('utf-8')
    
    # Remove code blocks
    content = re.sub(r'```[\s\S]*?```', '', content)
    content = re.sub(r'`[^`]+`', '', content)
    
    # Remove headers markers but keep text
    content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
    
    # Remove bold/italic markers
    content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)
    content = re.sub(r'\*([^*]+)\*', r'\1', content)
    content = re.sub(r'__([^_]+)__', r'\1', content)
    content = re.sub(r'_([^_]+)_', r'\1', content)
    
    # Remove links but keep text
    content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
    
    # Remove images
    content = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', content)
    
    # Remove horizontal rules
    content = re.sub(r'^[-*_]{3,}\s*$', '', content, flags=re.MULTILINE)
    
    # Remove list markers
    content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)
    content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)
    
    return content


def extract_text_from_html(file) -> str:
    """Extract text from HTML file, removing tags"""
    content = file.file.read().decode('utf-8')
    
    # Remove script and style elements
    content = re.sub(r'<script[^>]*>[\s\S]*?</script>', '', content, flags=re.IGNORECASE)
    content = re.sub(r'<style[^>]*>[\s\S]*?</style>', '', content, flags=re.IGNORECASE)
    
    # Remove all HTML tags
    content = re.sub(r'<[^>]+>', ' ', content)
    
    # Decode HTML entities
    import html
    content = html.unescape(content)
    
    # Clean up whitespace
    content = re.sub(r'\s+', ' ', content)
    
    return content.strip()


def extract_text(file) -> str:
    """
    Extract text from uploaded file
    
    Supported formats: PDF, TXT, DOCX, MD, HTML
    
    Args:
        file: FastAPI UploadFile object
        
    Returns:
        Extracted text content
    """
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.txt'):
        return extract_text_from_txt(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file)
    elif filename.endswith('.md') or filename.endswith('.markdown'):
        return extract_text_from_markdown(file)
    elif filename.endswith('.html') or filename.endswith('.htm'):
        return extract_text_from_html(file)
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. "
            f"Supported formats: PDF, TXT, DOCX, MD, HTML"
        )


def clean_text(text: str) -> str:
    """
    Clean and preprocess text
    
    - Remove excessive whitespace
    - Remove special characters that don't add meaning
    - Normalize unicode
    """
    # Normalize unicode
    import unicodedata
    text = unicodedata.normalize('NFKC', text)
    
    # Replace multiple whitespace with single space
    text = re.sub(r'\s+', ' ', text)
    
    # Remove very long sequences of special characters
    text = re.sub(r'[^\w\s]{5,}', ' ', text)
    
    return text.strip()


def split_text(text: str, chunk_size: int = 250, overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks
    
    Args:
        text: Text to split
        chunk_size: Target number of words per chunk
        overlap: Number of overlapping words between chunks
        
    Returns:
        List of text chunks
    """
    # Clean text first
    text = clean_text(text)
    
    words = text.split()
    
    if len(words) <= chunk_size:
        return [text] if text else []
    
    chunks = []
    start = 0
    
    while start < len(words):
        end = start + chunk_size
        chunk = ' '.join(words[start:end])
        
        if chunk.strip():
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - overlap
        
        # Prevent infinite loop for small texts
        if start >= len(words) - overlap:
            break
    
    return chunks


def split_text_by_sentences(text: str, max_chunk_size: int = 500) -> List[str]:
    """
    Split text by sentences, keeping chunks under max size
    
    This method preserves sentence boundaries for better context
    
    Args:
        text: Text to split
        max_chunk_size: Maximum characters per chunk
        
    Returns:
        List of text chunks
    """
    # Clean text
    text = clean_text(text)
    
    # Split by sentence endings
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chunk_size:
            current_chunk += " " + sentence if current_chunk else sentence
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks


def get_embeddings(texts: List[str]) -> List[List[float]]:
    """
    Generate embeddings for a list of texts
    
    Args:
        texts: List of text strings
        
    Returns:
        List of embedding vectors
    """
    embeddings = model.encode(texts)
    return embeddings.tolist()


def index_upload(file, chunk_size: int = 250, overlap: int = 50) -> int:
    """
    Process uploaded file: extract text, split, embed and store in vector database
    
    Args:
        file: FastAPI UploadFile object
        chunk_size: Words per chunk
        overlap: Overlap between chunks
        
    Returns:
        Number of chunks added
    """
    # Reset file pointer to beginning
    file.file.seek(0)
    
    # Extract text from file
    text = extract_text(file)
    
    if not text.strip():
        raise ValueError("No text content could be extracted from the file")
    
    # Split text into chunks
    chunks = split_text(text, chunk_size, overlap)
    
    if not chunks:
        raise ValueError("No chunks were generated from the text")
    
    # Generate embeddings
    embeddings = model.encode(chunks)
    
    # Store in vector database
    vector_db.add_to_vector_db(chunks, embeddings, file.filename)
    
    return len(chunks)


def index_text(text: str, source: str, chunk_size: int = 250, overlap: int = 50) -> int:
    """
    Process raw text: split, embed and store in vector database
    
    Args:
        text: Raw text content
        source: Source identifier
        chunk_size: Words per chunk
        overlap: Overlap between chunks
        
    Returns:
        Number of chunks added
    """
    if not text.strip():
        raise ValueError("Empty text provided")
    
    # Split text into chunks
    chunks = split_text(text, chunk_size, overlap)
    
    if not chunks:
        raise ValueError("No chunks were generated from the text")
    
    # Generate embeddings
    embeddings = model.encode(chunks)
    
    # Store in vector database
    vector_db.add_to_vector_db(chunks, embeddings, source)
    
    return len(chunks)


# Supported file extensions
SUPPORTED_EXTENSIONS = ['.pdf', '.txt', '.docx', '.md', '.markdown', '.html', '.htm']


def get_supported_formats() -> List[str]:
    """Return list of supported file formats"""
    return SUPPORTED_EXTENSIONS.copy()
